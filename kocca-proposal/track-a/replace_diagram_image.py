#!/usr/bin/env python3
"""Table[30]의 기존 파이프라인 이미지를 새 pipeline_diagram_v3.png로 교체.

python-docx의 이미지 교체: 기존 이미지 relationship의 바이너리를 새 이미지로 덮어쓰기.
"""

import shutil
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

SRC = Path("/home/ppak/SynologyDrive/ykpark/wizdata/붙임2. 신청양식_26제작지원_진입형"
           "/2. 신청양식_26제작지원_진입형_실증(플랫폼설루션)/1. 사업수행(필수)/1-1.docx")
NEW_IMG = Path("/home/ppak/clawd/kocca-proposal/track-a/pipeline_diagram_v3.png")


def main():
    # 백업
    ts = datetime.now().strftime('%Y%m%d_%H%M')
    bak = SRC.with_name(f'1-1_bak_{ts}.docx')
    shutil.copy2(SRC, bak)
    print(f"백업: {bak}")

    doc = Document(str(SRC))
    t = doc.tables[30]
    cell = t.cell(0, 0)

    # drawing 요소에서 이미지 rId 찾기
    ns = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    }

    replaced = False
    for para in cell.paragraphs:
        for run in para.runs:
            drawings = run._element.findall('.//wp:inline', ns) + run._element.findall('.//wp:anchor', ns)
            for drawing in drawings:
                blips = drawing.findall('.//a:blip', ns)
                for blip in blips:
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id:
                        # 해당 relationship의 이미지 파트를 교체
                        part = para.part
                        rel = part.rels[embed_id]
                        image_part = rel.target_part

                        # 새 이미지 바이너리로 교체
                        new_blob = NEW_IMG.read_bytes()
                        image_part._blob = new_blob

                        print(f"✅ 이미지 교체 완료: rId={embed_id}")
                        print(f"   원본 파트: {image_part.partname}")
                        print(f"   새 이미지: {NEW_IMG.name} ({len(new_blob):,} bytes)")
                        replaced = True

    if not replaced:
        print("⚠️ Table[30]에서 이미지를 찾지 못함")
        return

    doc.save(str(SRC))
    print(f"\n📄 저장: {SRC}")
    print(f"📦 백업: {bak}")


if __name__ == '__main__':
    main()
