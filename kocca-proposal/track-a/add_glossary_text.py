#!/usr/bin/env python3
"""Table[30] 도식화 셀에: 이미지 교체 + 아래에 용어 해설 텍스트 추가.

이미지는 기존 것을 바이너리 교체, 텍스트는 이미지 아래 paragraph로 추가.
"""

import shutil
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path("/home/ppak/SynologyDrive/ykpark/wizdata/붙임2. 신청양식_26제작지원_진입형"
           "/2. 신청양식_26제작지원_진입형_실증(플랫폼설루션)/1. 사업수행(필수)/1-1.docx")
NEW_IMG = Path("/home/ppak/clawd/kocca-proposal/track-a/pipeline_diagram_v3.png")

GLOSSARY = [
    ("LoRA (Low-Rank Adaptation)",
     "대규모 AI 모델 전체를 재학습하지 않고, 소량의 데이터(작가 원화 50~100장)만으로 "
     "특정 화풍을 학습시키는 경량 미세조정 기술. 학습 시간이 짧고(~40분/작가) GPU 메모리 사용이 적어 "
     "효율적이다. 'fuse'는 학습된 LoRA 가중치를 AI 모델 본체에 직접 합체시키는 방식으로, "
     "별도 어댑터 없이 추론 속도를 유지하면서 화풍을 적용할 수 있다."),

    ("ControlNet (Canny Edge Detection)",
     "원본 사진에서 윤곽선(엣지, 경계선)만 추출하여 AI에게 '이 구도와 윤곽대로 그려라'고 "
     "안내하는 조건부 생성 기술. Canny는 컴퓨터 비전에서 가장 널리 쓰이는 윤곽선 검출 알고리즘이다. "
     "원본 사진의 픽셀(색상, 질감)은 전달하지 않고 윤곽 정보만 전달하므로, "
     "AI가 학습된 화풍을 자유롭게 표현하면서도 원래 사진의 구도(인물 위치, 포즈, 배경 배치)를 유지할 수 있다. "
     "scale 0.65는 윤곽 가이드를 65% 강도로 반영한다는 의미이다."),

    ("IP-Adapter-FaceID (InsightFace)",
     "AI가 생성하는 이미지에서 특정 인물의 얼굴을 보존하는 기술. "
     "InsightFace라는 얼굴 인식 엔진이 원본 사진에서 눈·코·입의 비율, 얼굴형, 이목구비 특징 등을 "
     "512개 수치(512차원 임베딩)로 변환하고, 이를 IP-Adapter-FaceID가 이미지 생성 과정에 주입한다. "
     "그 결과 화풍이 만화 스타일로 완전히 바뀌어도 '같은 사람'임을 알아볼 수 있다. "
     "scale 0.40은 얼굴 보존을 40% 강도로 적용한다는 뜻으로, "
     "이보다 높으면 화풍 재현이 약해지고, 낮으면 얼굴이 달라지는 트레이드오프가 있다."),

    ("txt2img (Text-to-Image)",
     "텍스트 설명(프롬프트)만으로 이미지를 처음부터 생성하는 방식. "
     "이와 대비되는 img2img(Image-to-Image)는 원본 사진을 직접 입력으로 사용하는데, "
     "이 경우 원본 사진의 색상·질감이 결과물에 남아(픽셀 잔류) 화풍 재현을 방해한다. "
     "본 파이프라인은 txt2img를 채택하여 원본 픽셀의 간섭을 원천 차단하고, "
     "필요한 정보(윤곽, 얼굴)만 별도 기술(ControlNet, FaceID)로 전달함으로써 "
     "학습된 작가 화풍이 최대한 순수하게 반영되도록 설계하였다."),

    ("SDXL (Stable Diffusion XL)",
     "Stability AI가 개발한 오픈소스 이미지 생성 AI 모델. "
     "1024×1024 고해상도 이미지를 생성할 수 있으며, 텍스트 이해력과 이미지 품질이 우수하다. "
     "본 프로젝트의 기반 모델로, LoRA(화풍), ControlNet(구조), FaceID(얼굴) 등 "
     "추가 기술들과 결합하여 작가 화풍 변환 엔진의 핵심 역할을 수행한다."),

    ("guidance / steps / scale",
     "AI 이미지 생성의 품질을 조절하는 주요 파라미터. "
     "guidance(가이던스, 본 시스템에서 8.0)는 텍스트 프롬프트를 얼마나 충실히 따를지를 결정하며, "
     "높을수록 프롬프트에 충실하지만 다양성이 줄어든다. "
     "steps(스텝, 본 시스템에서 30)는 이미지를 만드는 반복 횟수로, 많을수록 정교하지만 시간이 늘어난다. "
     "scale은 ControlNet(0.65)과 FaceID(0.40) 등 보조 기술의 적용 강도(0~1)를 의미한다."),
]


def main():
    ts = datetime.now().strftime('%Y%m%d_%H%M')
    bak = SRC.with_name(f'1-1_bak_{ts}.docx')
    shutil.copy2(SRC, bak)
    print(f"백업: {bak}")

    doc = Document(str(SRC))

    # ── 1) 이미지 교체 ──
    t = doc.tables[30]
    cell = t.cell(0, 0)
    ns = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    }
    for para in cell.paragraphs:
        for run in para.runs:
            drawings = (run._element.findall('.//wp:inline', ns) +
                        run._element.findall('.//wp:anchor', ns))
            for drawing in drawings:
                for blip in drawing.findall('.//a:blip', ns):
                    embed_id = blip.get(f'{{{ns["r"]}}}embed')
                    if embed_id:
                        rel = para.part.rels[embed_id]
                        rel.target_part._blob = NEW_IMG.read_bytes()
                        print(f"  이미지 교체: rId={embed_id}")

    # ── 2) 기존 용어 해설 텍스트 제거 (있으면) ──
    # "용어 해설" 이후 paragraphs 제거
    found_glossary = False
    to_remove = []
    for para in cell.paragraphs:
        if '용어 해설' in para.text:
            found_glossary = True
            to_remove.append(para)
            continue
        if found_glossary:
            to_remove.append(para)

    for p in to_remove:
        p._element.getparent().remove(p._element)
        print(f"  기존 텍스트 제거: {p.text[:40]}...")

    # ── 3) 용어 해설 텍스트 추가 ──
    # 빈줄
    blank = cell.add_paragraph()
    blank.space_after = Pt(6)

    # 제목
    title_p = cell.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run('■ 용어 해설')
    title_run.font.size = Pt(11)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # 각 용어
    for term, desc in GLOSSARY:
        p = cell.add_paragraph()
        p.space_before = Pt(4)
        p.space_after = Pt(2)

        # 용어명 (bold)
        term_run = p.add_run(f'• {term}: ')
        term_run.font.size = Pt(10)
        term_run.font.bold = True
        term_run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)  # indigo

        # 설명
        desc_run = p.add_run(desc)
        desc_run.font.size = Pt(10)
        desc_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    doc.save(str(SRC))
    print(f"\n✅ 완료: 이미지 교체 + 용어 해설 텍스트 {len(GLOSSARY)}개 추가")
    print(f"📄 {SRC}")


if __name__ == '__main__':
    main()
