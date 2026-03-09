#!/usr/bin/env python3
"""DOCX 업데이트 스크립트 v3 — txt2img + ControlNet + FaceID 파이프라인 반영.

대상: 1-1.docx (SynologyDrive 원본)
"""

import shutil
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt


SRC = Path("/home/ppak/SynologyDrive/ykpark/wizdata/붙임2. 신청양식_26제작지원_진입형"
           "/2. 신청양식_26제작지원_진입형_실증(플랫폼설루션)/1. 사업수행(필수)/1-1.docx")


def set_cell(table, row, col, text: str):
    """셀 텍스트 교체 — 기존 서식 최대한 유지."""
    cell = table.cell(row, col)
    # 첫 번째 paragraph의 style/run format 보존
    paras = cell.paragraphs
    if not paras:
        return

    # 기존 스타일 캡쳐
    first_para = paras[0]
    para_style = first_para.style
    run_font_name = None
    if first_para.runs:
        r0 = first_para.runs[0]
        run_font_name = r0.font.name

    # 기존 paragraph 모두 제거 (첫 번째 제외)
    for p in paras[1:]:
        p_elem = p._element
        p_elem.getparent().remove(p_elem)

    # 첫 번째 paragraph 비우기
    for run in first_para.runs:
        run._element.getparent().remove(run._element)

    # 새 텍스트를 줄별로 작성
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            run = first_para.add_run(line)
        else:
            # 새 paragraph 추가
            new_p = cell.add_paragraph()
            new_p.style = para_style
            run = new_p.add_run(line)

        # 서식 적용 — 12pt 고정
        if run_font_name:
            run.font.name = run_font_name
        run.font.size = Pt(12)


def main():
    # 백업
    ts = datetime.now().strftime('%Y%m%d_%H%M')
    bak = SRC.with_name(f'1-1_bak_{ts}.docx')
    shutil.copy2(SRC, bak)
    print(f"백업: {bak}")

    doc = Document(str(SRC))
    updated = []

    # ═══════════════════════════════════════════════
    # Table[25]: 2-1 과제 개요 (8행 × 3열, col 1-2 merged)
    # ═══════════════════════════════════════════════
    t = doc.tables[25]

    # [1] 사업목적 및 기획 방향
    set_cell(t, 1, 2, (
        "본 과제는 7080 만화작가 IP(이정문, 신문수 등)의 고유 화풍을 AI로 학습·재현하여 "
        "사진 기반 콘텐츠 제작을 자동화하는 플랫폼을 실증하는 데 목적이 있다.\n"
        "자체 개발 AI 파이프라인(txt2img + ControlNet(Canny) + LoRA(fuse) + IP-Adapter-FaceID)을 통해 "
        "원본 사진의 구조와 얼굴 정체성을 보존하면서 작가 고유 화풍을 정밀 재현하며, "
        "이를 웹 SaaS로 구현하여 실제 콘텐츠 제작 현장에서 사업성을 검증한다.\n"
        "웹 기반 프로토타입 구축·운영 중(2026.02). 68건 이상 변환 실적, ~4초/장 변환 속도로 "
        "기술적 실현 가능성 확인 완료. 본 과제를 통해 상용 수준의 SaaS 플랫폼으로 고도화한다."
    ))

    # [2] 과제 주요내용
    set_cell(t, 2, 2, (
        "7080 만화작가 화풍 학습 모델을 구축하고, 사용자가 사진 업로드→작가/버전 선택→변환까지 "
        "수행 가능한 웹 기반 AI 플랫폼(ToonStyle AI)을 개발·실증한다.\n"
        "- txt2img 기반 단일 패스 파이프라인으로 화풍 재현(LoRA fuse) + 구조 보존(ControlNet Canny) "
        "+ 얼굴 정체성 유지(IP-Adapter-FaceID)를 동시 처리. "
        "작가 5인 이상 모델 확보 및 실증 3건 이상 수행을 목표로 한다.\n"
        "- PoC 실적: 2인 작가 모델(이정문·신문수) + LoRA 다중 버전(v1/v2/v2-strong) "
        "+ 웹서비스 운영 + 68건 이상 변환 테스트 완료."
    ))

    # [3] 차별성 — 기술수준 (col 2)
    set_cell(t, 3, 2, (
        "※ 기술 구현 범위, 자동화 수준 등을 기준으로 비교\n"
        "자체 개발 txt2img + ControlNet + FaceID 통합 파이프라인으로 특정 작가 화풍을 "
        "50~100장 원화만으로 정밀 재현. 범용 AI(Midjourney 등)는 특정 작가 화풍 학습 불가.\n"
        "- img2img 방식의 한계(원본 픽셀 잔류로 스타일 저하)를 극복한 txt2img 설계: "
        "Canny 엣지만 구조 가이드로 전달하여 LoRA 화풍 자유도 극대화.\n"
        "- InsightFace 512d 얼굴 임베딩 + IP-Adapter-FaceID로 화풍 변환 후에도 인물 식별 가능"
        "(화풍 9/10, 얼굴 7/10).\n"
        "- 정식 IP 라이선싱으로 저작권 리스크 제로. LoRA 다중 버전(v1/v2/v2-strong)으로 화풍 강도 선택 가능."
    ))

    # [4] 제작·활용 방식
    set_cell(t, 4, 2, (
        "※ 예) 어떤 콘텐츠를 어떠한 기술로 어디에서 활용할 것인지 간략하게 작성\n"
        "웹에서 사진 업로드 → 작가/LoRA 버전 선택 → AI 변환(~4초, 단일 패스) → "
        "결과 다운로드/갤러리 저장. 프리셋 5종으로 최적 파라미터 즉시 적용 가능.\n"
        "- 파라미터(FaceID 강도, ControlNet 강도, 가이던스 등) 직접 조절 UI 제공. "
        "기존 만화 작가 직접 작화 대비 시간 99.9% 단축(수 시간→~4초). "
        "B2B API 연동으로 외부 서비스 자동화 가능."
    ))

    # [5] 활용 범위
    set_cell(t, 5, 2, (
        "※ 플랫폼/설루션 타겟 소비자, 이용자, 적용 범위 및 적용 분야 등\n"
        "SNS/유튜브 크리에이터 콘텐츠, 기업 마케팅·브랜딩, 전시·교육 체험형 콘텐츠, "
        "출판/디자인 시안 제작 등으로 확장 가능.\n"
        "- 작가 IP 추가에 따라 활용 분야를 지속 확대. "
        "LoRA 다중 버전(기본/밸런스/최대 화풍)으로 용도별 최적 결과 제공."
    ))

    # [6] 사업성
    set_cell(t, 6, 2, (
        "B2C 구독형(월 9,900~29,900원) + B2B API 라이선스(월 100만~500만원) "
        "+ IP 라이선싱 연계의 복합 수익모델 적용.\n"
        "- 글로벌 AI 이미지 생성 시장 2028년 약 80억 달러 전망(CAGR 52%). "
        "K-콘텐츠 글로벌 확산 흐름과 레트로 콘텐츠 수요 기반, "
        "한국 고유 만화 IP 특화 서비스로 차별화된 시장 진입 가능."
    ))

    # [7] 기대효과
    set_cell(t, 7, 2, (
        "① 한국 고유 만화 IP의 디지털 전환 및 재사업화 촉진\n"
        "② AI 기반 콘텐츠 제작 생태계 확장, 창작자 친화적 생성형 AI 활용 모델 제시\n"
        "③ 제작시간·비용 절감에 따른 창작 효율 향상(수 시간→~4초/장, 단일 패스)\n"
        "④ 신규 일자리 창출(AI 엔지니어, 콘텐츠 기획) 및 후속 투자 연계"
    ))
    updated.append("Table[25] 과제 개요")

    # ═══════════════════════════════════════════════
    # Table[27]: 세부 제작계획 (5행 × 4열, col 2-3 merged)
    # ═══════════════════════════════════════════════
    t = doc.tables[27]

    # [0] 플랫폼/솔루션
    set_cell(t, 0, 2, (
        "ㅇ ToonStyle AI 웹 플랫폼 — 프로토타입 운영 중(2026.02). "
        "사진 업로드→작가/버전 선택→단일 패스 AI 변환(txt2img + ControlNet + FaceID)"
        "→갤러리 관리까지 웹 UI 동작 확인.\n"
        "- 현재 2인 작가 모델(이정문·신문수) + LoRA 3종 버전(v1/v2/v2-strong) 탑재. "
        "본 과제에서 5인 이상으로 확장.\n"
        "- FastAPI 백엔드 + Tailwind CSS 프론트엔드 + GPU 추론 인프라. "
        "프리셋 5종, FaceID/ControlNet 강도 슬라이더, 변환 이력 갤러리, B2B API. "
        "변환 속도 ~4초/장(1024px)."
    ))

    # [1] 실증 콘텐츠
    set_cell(t, 1, 2, (
        "ㅇ 실증 3트랙:\n"
        "① SNS/유튜브 크리에이터 협업 화풍 변환 콘텐츠 시리즈\n"
        "② 기업 마케팅용 캐릭터 굿즈·프로필·홍보 이미지\n"
        "③ 7080 만화 전시·체험 콘텐츠\n\n"
        "ㅇ PoC 단계에서 이미 68건 이상 변환 테스트 완료. "
        "실증기간 중 10,000건 이상 생성, 분야별 적용사례 리포트 작성."
    ))

    # [2] 콘텐츠 품질 개선 요소
    set_cell(t, 2, 2, (
        "ㅇ txt2img 기반 설계로 원본 픽셀 잔류 없이 순수 화풍 생성 — "
        "img2img 대비 스타일 재현도 대폭 향상(화풍 점수 9/10 달성).\n"
        "- ControlNet(Canny)으로 원본 구조(윤곽·포즈) 보존 + InsightFace FaceID로 "
        "얼굴 정체성 유지. 단일 패스로 화풍·구조·얼굴 동시 처리.\n"
        "- LoRA 다중 버전(v1 기본/v2 밸런스/v2-strong 최대 화풍) + "
        "프리셋 시스템으로 품질 일관성 유지."
    ))
    updated.append("Table[27] 세부 제작계획")

    # ═══════════════════════════════════════════════
    # Table[28]: AI 기술 활용 계획 (2행 × 2열)
    # ═══════════════════════════════════════════════
    t = doc.tables[28]

    # [0] AI 기술 설명
    set_cell(t, 0, 1, (
        "ㅇ 자체 개발 단일 패스 AI 화풍 변환 파이프라인을 핵심 엔진으로 구축. "
        "txt2img 방식으로 원본 사진 픽셀을 직접 사용하지 않고, "
        "구조(Canny 엣지)와 얼굴(FaceID 임베딩)만 추출하여 가이드로 전달함으로써 "
        "LoRA 화풍 재현도를 극대화하는 설계.\n\n"
        "- txt2img + LoRA(fuse): 작가 화풍을 학습한 LoRA 가중치를 SDXL UNet에 직접 fuse하여 "
        "텍스트 프롬프트 기반으로 이미지 생성. 원본 픽셀에 의한 스타일 간섭 제거.\n"
        "- ControlNet(Canny): 원본 사진에서 Canny 엣지맵을 추출하여 구도·윤곽 가이드로 전달"
        "(scale=0.65). 사진 구조 유지와 화풍 자유도의 최적 밸런스.\n"
        "- IP-Adapter-FaceID: InsightFace(buffalo_l)로 512차원 얼굴 임베딩 추출 → "
        "IP-Adapter-FaceID로 생성 과정에 주입(scale=0.40). 화풍 변환 후에도 인물 식별 가능.\n"
        "- LoRA 버전 시스템: 작가별 v1(기본)/v2(밸런스·epoch-6)/v2-strong(최대 화풍·epoch-8) 다중 버전 제공.\n\n"
        "- 작가 화풍 학습: 50~100장 원화로 SDXL LoRA 학습(dim=32, alpha=16). "
        "이정문·신문수 2인 학습 완료, 본 과제에서 5인 이상 확장."
    ))

    # [1] AI 활용 효과
    set_cell(t, 1, 1, (
        "ㅇ 기존 만화 작가 직접 작화 시 1컷당 수 시간~수일 소요 → AI 변환으로 ~4초/장으로 단축"
        "(제작시간 99.9% 이상 절감).\n"
        "- 기존 범용 AI(Midjourney, ChatGPT 등)로는 특정 작가 고유 화풍을 재현할 수 없음. "
        "본 기술은 50~100장 원화만으로 작가별 화풍을 정밀 학습하여, "
        "정식 라이선스 기반의 작가 특화 변환을 가능하게 함.\n"
        "- 기존 img2img 방식의 한계(원본 픽셀 잔류로 화풍 재현 저하)를 "
        "txt2img + ControlNet 설계로 근본 해결 — LoRA 스타일 재현도 9/10 달성.\n"
        "- InsightFace FaceID 기반 얼굴 보존으로 화풍 변환 후에도 인물 식별 가능"
        "(얼굴 보존 7/10). 기존 IP-Adapter 스타일 참조 방식 대비 얼굴 특화 성능 향상."
    ))
    updated.append("Table[28] AI 기술 활용 계획")

    # ═══════════════════════════════════════════════
    # Table[29]: (4-1) AI 기술 기능 상세 (1행 × 1열)
    # ═══════════════════════════════════════════════
    t = doc.tables[29]
    set_cell(t, 0, 0, (
        "ㅇ 자체 개발 단일 패스 txt2img + ControlNet + FaceID 통합 파이프라인을 핵심 엔진으로, "
        "원본 사진 픽셀을 사용하지 않는 설계로 LoRA 화풍 재현도를 극대화하면서 "
        "구조·얼굴 보존을 동시 확보\n"
        "- 기존 작가 직접 작화(수 시간/건) → AI 변환(~4초/건)으로 콘텐츠 제작 시간 99.9% 단축. "
        "50~100장 원화로 작가 화풍 학습 가능\n\n"
        "기술 내역                          | 콘텐츠 제작 단계           | 활용 방식                         | 주요 내용\n"
        "───────────────────────────────────────────────────────────────────────────────────\n"
        "SDXL txt2img + LoRA(fuse)          | AI 모델 학습/화풍 변환     | 작가별 화풍 학습 및 txt2img 변환   | 원화 50~100장으로 학습. LoRA를 UNet에 fuse하여 화풍 자유 생성. 다중 버전(v1/v2/v2-strong)\n"
        "ControlNet(Canny)                  | 구조 가이드              | Canny 엣지맵으로 구도·윤곽 유지    | control_strength 0.65에서 최적 밸런스. 원본 픽셀 직접 사용 않음\n"
        "IP-Adapter-FaceID(InsightFace)     | 얼굴 정체성 보존          | 512d 얼굴 임베딩을 생성 과정에 주입 | faceid_scale 0.40에서 화풍 9/10 + 얼굴 7/10 동시 달성\n"
        "자동 캡셔닝 + 품질 평가             | 데이터 전처리/결과물 검증   | 이미지별 자동 분석 캡션 + 품질 평가 | 트리거 워드 방식. SSIM·LPIPS·FID 자동 평가"
    ))
    updated.append("Table[29] AI 기술 기능 상세")

    # ═══════════════════════════════════════════════
    # Table[30]: (4-2) AI 기술 도식화 (1행 × 1열)
    # ═══════════════════════════════════════════════
    t = doc.tables[30]
    set_cell(t, 0, 0, (
        "[추론 파이프라인 — 단일 패스 변환]\n\n"
        "[사용자 입력]\n"
        "  사진 업로드 / 화백·버전 선택 / 파라미터 조절\n"
        "       │\n"
        "       ▼\n"
        "┌─── AI 변환 파이프라인 (단일 패스, ~4초) ───┐\n"
        "│                                           │\n"
        "│  원본 사진 → Canny Edge 추출               │\n"
        "│       └→ ControlNet (scale=0.65)          │\n"
        "│           구도·윤곽 가이드                  │\n"
        "│                    ↓                      │\n"
        "│  원본 얼굴 → InsightFace (buffalo_l)       │\n"
        "│       └→ 512d 임베딩 → FaceID (scale=0.40)│\n"
        "│           얼굴 정체성 주입                  │\n"
        "│                    ↓                      │\n"
        "│        SDXL txt2img + LoRA (fuse)         │\n"
        "│        guidance=8.0 / steps=30            │\n"
        "│        → 화풍 변환 이미지 생성               │\n"
        "└───────────────────────────────────────────┘\n"
        "       │\n"
        "       ▼\n"
        "[출력] 화풍 변환 이미지 다운로드 / 갤러리 저장\n\n"
        "※ 핵심: 원본 사진 픽셀을 직접 사용하지 않음(img2img ✕).\n"
        "  Canny 엣지(구조)와 FaceID 임베딩(얼굴)만 추출하여\n"
        "  가이드로 전달 → LoRA 화풍 재현도 극대화.\n\n"
        "[모델 학습 파이프라인]\n"
        "작가 원화(50~100장) → 이미지별 자동 분석 캡셔닝\n"
        "(컬러/흑백·배경·구도 분석, 트리거 워드 삽입)\n"
        "→ SDXL LoRA 학습(dim=32, alpha=16, ~40분/작가)\n"
        "→ TensorBoard 모니터링 + epoch별 과적합 검출\n"
        "→ 작가 화풍 모델 다중 버전(v1/v2/v2-strong)"
    ))
    updated.append("Table[30] AI 기술 도식화")

    # ═══════════════════════════════════════════════
    # Save
    # ═══════════════════════════════════════════════
    doc.save(str(SRC))

    print(f"\n✅ DOCX v3 업데이트 완료: {SRC}")
    print(f"📦 백업: {bak}")
    print(f"📝 반영된 테이블:")
    for t in updated:
        print(f"   ✏️  {t}")


if __name__ == '__main__':
    main()
