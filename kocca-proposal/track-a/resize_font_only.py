#!/usr/bin/env python3
"""DOCX 폰트 크기만 변경 — 텍스트/이미지 등 내용은 건드리지 않음.

우리가 수정한 테이블(25, 27, 28, 29, 30)의 폰트를 12pt로 키움.
"""

import shutil
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt

SRC = Path("/home/ppak/SynologyDrive/ykpark/wizdata/붙임2. 신청양식_26제작지원_진입형"
           "/2. 신청양식_26제작지원_진입형_실증(플랫폼설루션)/1. 사업수행(필수)/1-1.docx")

TARGET_TABLES = [25, 27, 28, 29, 30]
NEW_SIZE = Pt(12)


def main():
    # 백업
    ts = datetime.now().strftime('%Y%m%d_%H%M')
    bak = SRC.with_name(f'1-1_bak_{ts}.docx')
    shutil.copy2(SRC, bak)
    print(f"백업: {bak}")

    doc = Document(str(SRC))

    changed = 0
    for tidx in TARGET_TABLES:
        t = doc.tables[tidx]
        for row in t.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.size != NEW_SIZE:
                            run.font.size = NEW_SIZE
                            changed += 1

    doc.save(str(SRC))
    print(f"\n✅ 폰트 크기 변경 완료: {changed}개 run → 12pt")
    print(f"📄 파일: {SRC}")
    print(f"📦 백업: {bak}")
    print(f"🎯 대상 테이블: {TARGET_TABLES}")


if __name__ == '__main__':
    main()
